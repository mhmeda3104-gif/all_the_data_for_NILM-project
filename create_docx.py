import docx
from docx.shared import Inches
import os
import re

def create_word_doc(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('!['):
            # Image
            match = re.search(r'\((.*?)\)', line)
            if match:
                img_path = match.group(1).strip()
                try:
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph(line.split('](')[0].replace('![', 'Figure: '))
                except Exception as e:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Saved {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\POWER\.gemini\antigravity\brain\dd368053-1cdf-47a1-be6b-60c13c433f93\NILM_Research_Paper.md"
    docx_file = r"C:\Users\POWER\Desktop\NILM\NILM_Research_Paper.docx"
    create_word_doc(md_file, docx_file)

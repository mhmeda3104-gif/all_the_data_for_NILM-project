from docx import Document
doc = Document(r"C:\Users\POWER\Desktop\NILM\NILM_Final_Paper.docx")
text = ""
for p in doc.paragraphs:
    text += p.text + "\n"
words = text.split()
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total words: {len(words)}")
print(f"Total characters: {len(text)}")

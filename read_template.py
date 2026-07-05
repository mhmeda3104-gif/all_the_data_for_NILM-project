import docx
from docx import Document
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document(r'c:\Users\POWER\Desktop\NILM\nurai2026template.docx')

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] Style='{para.style.name}' | Text: {para.text[:120]}")

print("\n\n--- TABLES ---")
for t, table in enumerate(doc.tables):
    print(f"\nTable {t}: {len(table.rows)} rows x {len(table.columns)} cols")
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  [{r},{c}]: {cell.text[:80]}")

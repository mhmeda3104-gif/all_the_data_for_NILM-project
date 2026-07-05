import docx
from docx import Document
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document(r'c:\Users\POWER\Desktop\NILM\nurai2026template.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
# Print the last 30 paragraphs of the template to see what is at the end of the document
for i in range(max(0, len(doc.paragraphs)-35), len(doc.paragraphs)):
    print(f"[{i}] {doc.paragraphs[i].text[:120]}")

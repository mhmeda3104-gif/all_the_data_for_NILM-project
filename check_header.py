from docx import Document
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document(r'c:\Users\POWER\Desktop\NILM\nurai2026template.docx')

for s, section in enumerate(doc.sections):
    print(f"\n--- Section {s} ---")
    header = section.header
    footer = section.footer
    print(f"Header paragraphs: {len(header.paragraphs)}")
    for p in header.paragraphs:
        if p.text.strip():
            print(f"  H: {p.text}")
    print(f"Footer paragraphs: {len(footer.paragraphs)}")
    for p in footer.paragraphs:
        if p.text.strip():
            print(f"  F: {p.text}")
